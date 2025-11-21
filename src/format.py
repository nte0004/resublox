from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement as SharedOxmlElement, qn as shared_qn
from docx.text.run import Run
from docx.enum.style import WD_STYLE_TYPE
import docx.opc.constants
import tempfile
import os
import platform
import subprocess
import zipfile
import shutil
from lxml import etree
from src.core import FontMetrics, PAGE_WIDTH_INCHES, PAGE_HEIGHT_INCHES, MARGIN_INCHES, FONT
from src.lineGenerator import LineGenerator
from src.linkHandler import LinkFormatter

FONT_METRICS = FontMetrics()
LINE_GENERATOR = LineGenerator(FONT_METRICS)
LINK_FORMATTER = LinkFormatter()

def get_or_create_hyperlink_style(document):
    """
    If this document had no hyperlinks so far, the builtin
    Hyperlink style will likely be missing and we need to add it.
    
    Based on Stack Overflow solution by planet260
    Source: https://stackoverflow.com/a/47666747
    License: CC BY-SA 4.0
    """
    if "Hyperlink" not in document.styles:
        if "Default Character Font" not in document.styles:
            ds = document.styles.add_style("Default Character Font",
                                    WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(shared_qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = document.styles.add_style("Hyperlink",
                                WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = document.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs
    return "Hyperlink"

def add_hyperlink_to_paragraph(paragraph, text, url, font_name=None, font_size=None):
    """
    Add a hyperlink to a paragraph with proper formatting.
    
    Based on Stack Overflow solution by planet260
    Source: https://stackoverflow.com/a/47666747
    License: CC BY-SA 4.0
    Modified to support custom font name and size
    
    :param paragraph: The paragraph to add the hyperlink to
    :param text: The text to display
    :param url: The URL to link to
    :param font_name: Optional font name
    :param font_size: Optional font size (Pt object)
    :return: The hyperlink element
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Create the w:hyperlink tag and add needed values
    hyperlink = SharedOxmlElement('w:hyperlink')
    hyperlink.set(shared_qn('r:id'), r_id)
    
    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = Run(SharedOxmlElement('w:r'), paragraph)
    new_run.text = text
    
    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)
    
    # Apply custom font formatting if specified
    if font_name:
        new_run.font.name = font_name
    
    if font_size:
        new_run.font.size = font_size
    
    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def add_line(paragraph, lineSpec):
    """
    Add a line to the paragraph, handling links if present.
    
    :param paragraph: The paragraph to add content to
    :param lineSpec: The LineSpec containing text and optional links
    """
    font_name = FONT.name
    font_size = Pt(lineSpec.size.size)
    
    if lineSpec.links and len(lineSpec.links) > 0:
        # Process links using the LinkFormatter
        formatted_links = LINK_FORMATTER.format_collection_for_docx(lineSpec.links)
        
        for _, (prefix, display, url) in enumerate(formatted_links):
            if prefix and display and url:
                # This is a link with prefix
                run = paragraph.add_run(prefix)
                run.font.name = font_name
                run.font.size = font_size
                
                # Add the hyperlink
                try:
                    add_hyperlink_to_paragraph(paragraph, display, url, font_name, font_size)
                except Exception as e:
                    print(f"Warning: Failed to add hyperlink: {e}")
                    # Fallback to plain text
                    run = paragraph.add_run(display)
                    run.font.name = font_name
                    run.font.size = font_size
            elif prefix and not display and not url: # This is garbage I think.
                # This is just separator text
                run = paragraph.add_run(prefix)
                run.font.name = font_name
                run.font.size = font_size
            else:
                # Shouldn't happen, but handle gracefully
                if prefix:
                    run = paragraph.add_run(prefix)
                    run.font.name = font_name
                    run.font.size = font_size

        # NOTE: This is specifically so location in the contact section will work.
        #       not a huge fan of the solution but it works.
        links_display = lineSpec.links.get_display_text()
        if lineSpec.text and len(lineSpec.text) > len(links_display):
            # There's additional text (like " | City, State")
            additional_text = lineSpec.text[len(links_display):]
            run = paragraph.add_run(additional_text)
            run.font.name = font_name
            run.font.size = font_size
    else:
        # No links, just add regular text
        run = paragraph.add_run(lineSpec.text)
        run.font.name = font_name
        run.font.size = font_size

def generateLines(content):
    """Generate all lines for the resume"""
    lines = []
    
    lines.extend(LINE_GENERATOR.generateContactLines(content['contact']))
    lines.extend(LINE_GENERATOR.generateSkillsHeader(content['skills']))
    
    skillsContent = LINE_GENERATOR.generateSkillsContent(content['skills']['list'])
    lines.append(skillsContent)
    
    lines.extend(LINE_GENERATOR.generateExperienceHeader(content['experience']))
    
    for jobIdx, job in enumerate(content['experience']['jobs']):
        lines.extend(LINE_GENERATOR.generateJobHeader(job, jobIdx))
        
        for sectionIdx, section in enumerate(job['sections']):
            isFirstInJob = (sectionIdx == 0)
            lines.extend(LINE_GENERATOR.generateSectionHeader(section, jobIdx, sectionIdx, isFirstInJob))
            
            for pointIdx, point in enumerate(section['points']):
                lines.append(LINE_GENERATOR.generatePointLine(point, jobIdx, sectionIdx, pointIdx))
            
            if 'keywords' in section and section['keywords']:
                lines.append(LINE_GENERATOR.generateKeywordsLine(section['keywords'], jobIdx, sectionIdx))
            
            if 'links' in section and section['links'] is not None:
                lines.append(LINE_GENERATOR.generateLinksLine(section['links'], jobIdx, sectionIdx))
    
    if 'projects' in content and content['projects'] is not None:
        lines.extend(LINE_GENERATOR.generateProjectsHeader(content['projects']))
        
        for projIdx, project in enumerate(content['projects']['projects']):
            lines.extend(LINE_GENERATOR.generateProjectHeader(project, projIdx))
            
            for pointIdx, point in enumerate(project['points']):
                lines.append(LINE_GENERATOR.generateProjectPointLine(point, projIdx, pointIdx))
            
            if 'keywords' in project and project['keywords']:
                lines.append(LINE_GENERATOR.generateProjectKeywordsLine(project['keywords'], projIdx))
            
            if 'links' in project and project['links'] is not None:
                lines.append(LINE_GENERATOR.generateProjectLinksLine(project['links'], projIdx))
    
    lines.extend(LINE_GENERATOR.generateEducationLines(content['education']))

    if 'courses' in content['education'] and content['education']['courses']:
        coursesLine = LINE_GENERATOR.generateCoursesLine(content['education']['courses'])
        lines.append(coursesLine)
        
    return lines

def createDocx(lineSpecs):
    """Create a DOCX document from LineSpec objects"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_INCHES)
    section.page_height = Inches(PAGE_HEIGHT_INCHES)
    section.top_margin = Inches(MARGIN_INCHES[0])
    section.right_margin = Inches(MARGIN_INCHES[1])
    section.bottom_margin = Inches(MARGIN_INCHES[2])
    section.left_margin = Inches(MARGIN_INCHES[3])

    for lineSpec in lineSpecs:
        if lineSpec.text:
            insertion = doc.add_paragraph()
            
            add_line(insertion, lineSpec)
            
            insertion.paragraph_format.space_after = Pt(0)
            insertion.paragraph_format.space_before = Pt(0)
        else:
            # Empty line for spacing
            insertion = doc.add_paragraph()
            insertion.paragraph_format.space_after = Pt(lineSpec.size.size)
            insertion.paragraph_format.space_before = Pt(0)
            insertion.paragraph_format.line_spacing = Pt(0)

    return doc

def fix_hyperlinks_in_docx(docx_path):
    """
    Post-process a DOCX file to add w:history="1" to all hyperlinks.
    This is needed for LibreOffice PDF conversion to preserve hyperlinks.
    """
    # Create a temporary directory to extract the DOCX
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Extract the DOCX (which is a ZIP file)
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # Parse the document.xml
        doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        tree = etree.parse(doc_xml_path)
        root = tree.getroot()
        
        # Define namespaces
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        
        # Find all hyperlink elements and add w:history="1"
        hyperlinks = root.findall('.//w:hyperlink', namespaces)
        for hyperlink in hyperlinks:
            hyperlink.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}history', '1')
        
        # Write the modified XML back
        tree.write(doc_xml_path, xml_declaration=True, encoding='UTF-8', standalone=True)
        
        # Re-create the DOCX file
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as docx:
            for foldername, subfolders, filenames in os.walk(temp_dir):
                for filename in filenames:
                    file_path = os.path.join(foldername, filename)
                    arcname = os.path.relpath(file_path, temp_dir)
                    docx.write(file_path, arcname)
    
    finally:
        # Clean up temp directory
        shutil.rmtree(temp_dir)

def docxToPdf(tempPath):
    """Convert DOCX to PDF using LibreOffice"""
    system = platform.system()
    if system == "Darwin":
        LIBREOFFICE_PATH='/Applications/LibreOffice.app/Contents/MacOS/soffice'
    elif system == "Linux":
        LIBREOFFICE_PATH='libreoffice'
    else:
        print("Error: Windows docx to pdf conversion not implemented.")
        exit(1)

    subprocess.run([
        LIBREOFFICE_PATH, '--headless', '--convert-to', 'pdf',
        '--outdir', os.path.dirname(tempPath), tempPath
    ], check=True)

def output(content, editableFlag = False):
    """Main output function to generate the resume"""
    lineSpecs = generateLines(content)
    doc = createDocx(lineSpecs)

    # TODO: Named path here, passed from resublox.
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmpFile:
        doc.save(tmpFile.name)
        tempPath = tmpFile.name
    
    # Post-process to add w:history="1" to hyperlinks
    fix_hyperlinks_in_docx(tempPath)
    
    if (editableFlag):
        # Open as DOCX for editing
        subprocess.run(['open', tempPath])
    else:
        # Convert to PDF and open
        docxToPdf(tempPath)
        pdfPath = tempPath.replace('.docx', '.pdf')
        subprocess.run(['open', pdfPath])
