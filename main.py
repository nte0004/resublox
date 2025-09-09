import yaml
from docx import Document
from docx.shared import Inches, Pt
from fontTools.ttLib import TTFont
import math
import tempfile
import os
import platform
import subprocess
from dataclasses import dataclass

class FontSize:
    NAME = 13
    REGULAR = 10
    TITLE = 12
    SUBTITLE = 11

class Spacing:
    GAP = 6
    GAP_SMALL = 3

@dataclass
class FontInfo:
    name: str
    path: str
    unitsPerEm: int
    fontHeightUnits: int
    fontAvgWidthUnits: int

class Fonts:
    # Values found with fontHelper.py
    ARIAL = FontInfo(
        name = 'arial',
        path = 'fonts/arial/arial.ttf',
        unitsPerEm = 2048,
        fontHeightUnits = 2355,
        fontAvgWidthUnits = 1079
    )

TEMPLATE_PATH='template.example.yaml'
FONT = Fonts.ARIAL
PAGE_WIDTH_INCHES = 8.5
PAGE_HEIGHT_INCHES = 11
MARGIN_INCHES = [1, 1, 1, 1]
LINE_HEIGHT = 1.15
MAX_PAGES = 1

def combine(texts, seperator):
    line = ""
    for text in texts[:-1]:
        line += text
        line += seperator
    line += texts[-1];

    return line

def generateLines():
    lines = []
    
    with open(TEMPLATE_PATH, 'r') as file:
        content = yaml.safe_load(file)

    # Contact Info
    c = content['contact']
    lines.append((c['name'], FontSize.NAME))

    line = f"Email: {c['email']} | Phone: {c['phone']} | {c['location']}"
    lines.append((line, FontSize.REGULAR))

    line = f"Github: {c['github']} | Website: {c['website']}"
    lines.append((line, FontSize.REGULAR))

    # Gap
    lines.append(('', Spacing.GAP))

    # Skills
    s = content['skills']
    lines.append((s['title'], FontSize.TITLE))

    skills = combine(s['list'], ', ')
    lines.append((skills, FontSize.REGULAR))
    
    # Gap
    lines.append(('', Spacing.GAP))

    # Experience
    e = content['experience']
    lines.append((e['title'], FontSize.TITLE))
    for job in e['jobs']:
        lines.append((job['role'], FontSize.REGULAR))
        lines.append((job['company'], FontSize.REGULAR))
        lines.append((job['location'], FontSize.REGULAR))

        timeRange = combine([job['from'], job['to']], ' \u2014 ') # Em Dash
        lines.append((timeRange, FontSize.REGULAR))

        lines.append(('', Spacing.GAP_SMALL))

        sections = job['sections']
        lastSec = len(sections) - 1
        for idx, sec in enumerate(job['sections']):
            lines.append((sec['title'], FontSize.SUBTITLE))
            
            for point in sec['points']:
                lines.append((f"- {point}", FontSize.REGULAR))

            if "keywords" in sec:
                keywords = combine(sec['keywords'], ', ')
                lines.append((f"- Technologies: {keywords}", FontSize.REGULAR))

            if "links" in sec:
                links = []
                for link in sec['links']:
                    links.append(f"{link['descriptor']}: {link['link']}")
                linkLine = combine(links, ' | ');
                lines.append((linkLine, FontSize.REGULAR))
            
            if idx != lastSec:
                lines.append(('', Spacing.GAP_SMALL))

    # Gap
    lines.append(('', Spacing.GAP))
    
    # Education
    e = content['education']
    lines.append((e['title'], FontSize.TITLE))
    
    degree = f"{e['degree']} in {e['major']}"
    lines.append((degree, FontSize.REGULAR))

    conc = e.get('concentration', None)
    if conc:
        lines.append((f"Concentration in {conc}", FontSize.REGULAR))

    school = f"{e['school']}, {e['location']}"
    lines.append((school, FontSize.REGULAR))

    grad = e.get('graduation', None)
    if grad:
        print(grad)
        if grad['hasGraduated']:
            gradLine = f"Graduated: {grad['on']}"
        else:
            gradLine = f"Expected Graduation: {grad['on']}"
        lines.append((gradLine, FontSize.REGULAR))

    gpa = f"GPA: {e['gpa']}"
    lines.append((gpa, FontSize.REGULAR))

    honorsList = combine(e['honors'], ', ')
    honors = f"Honors: {honorsList}"
    lines.append((honors, FontSize.REGULAR))

    courseList = combine(e['courses'], ', ')
    courses = f"Relevant Courses: {courseList}"
    lines.append((courses, FontSize.REGULAR))

    return lines

def findOverflows(lines, warn = False):
    font = TTFont(FONT.path)
    cmap = font.getBestCmap()
    hmtx = font['hmtx']
    maxWidthInches = PAGE_WIDTH_INCHES - MARGIN_INCHES[1] - MARGIN_INCHES[3]

    maxHeightInches = (PAGE_HEIGHT_INCHES - MARGIN_INCHES[0] - MARGIN_INCHES[2]) * MAX_PAGES
    totalHeightInches = 0

    lineOverflows = []
    pageOverflows = []
    for text, size in lines:
        if text:
            totalWidth = 0
            for char in text:
                glyphName = cmap.get(ord(char))
                if glyphName:
                    width, _ = hmtx[glyphName]
                    totalWidth += width
                else:
                    totalWidth += FONT.fontAvgWidthUnits
            
            widthPts = (totalWidth * size) / FONT.unitsPerEm
            widthInches = widthPts / 72

            if widthInches > maxWidthInches:
                lineOverflows.append(text)
                if warn:
                    print(f"WARNING: This line will wrap: '{text}'")

            heightPts = (FONT.fontHeightUnits * size * LINE_HEIGHT) / FONT.unitsPerEm
            heightInches = heightPts / 72

            lines = math.ceil(widthInches / maxWidthInches) or 1

            totalHeightInches += heightInches * lines

            if totalHeightInches > maxHeightInches:
                pageOverflows.append(text)
                if warn:
                    print(f"WARNING: This line will overflow into a new page: '{text}'")
        else:
            spaceInches = size / 72
            totalHeightInches += spaceInches
            
            if totalHeightInches > maxHeightInches:
                pageOverflows.append(text)
                if warn:
                    print(f"WARNING: This line will overflow into a new page: '{text}'")

    return lineOverflows, pageOverflows

def createDocx(lines):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_INCHES)
    section.page_height = Inches(PAGE_HEIGHT_INCHES)
    section.top_margin = Inches(MARGIN_INCHES[0])
    section.right_margin = Inches(MARGIN_INCHES[1])
    section.bottom_margin = Inches(MARGIN_INCHES[2])
    section.left_margin = Inches(MARGIN_INCHES[3])

    for text, size in lines:
        if text:
            insertion = doc.add_paragraph()
            run = insertion.add_run(text)
            run.font.name = FONT.name
            run.font.size = Pt(size)

            insertion.paragraph_format.space_after = Pt(0)
            insertion.paragraph_format.space_before = Pt(0)
        else:
            insertion = doc.add_paragraph()
            insertion.paragraph_format.space_after = Pt(size)
            insertion.paragraph_format.space_before = Pt(0)
            insertion.paragraph_format.line_spacing = Pt(0)

    return doc

def docxToPdf(doc):
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmpFile:
        doc.save(tmpFile.name)
        tempPath = tmpFile.name

    system = platform.system()
    if system == "Darwin":
        LIBREOFFICE_PATH='/Applications/LibreOffice.app/Contents/MacOS/soffice'
    elif system == "Linux":
        LIBREOFFICE_PATH='libreoffice'
    else:
        print("Error: Windows docx to pdf conversion not implemented.")
        exit(1)

    subprocess.run([
        LIBREOFFICE_PATH, '--headless', '--convert-to', 'pdf',
        '--outdir', os.path.dirname(tempPath), tempPath
    ])

    return tempPath

def openPdf(path):
    pdfPath = path.replace('.docx', '.pdf')
    subprocess.run(['open', pdfPath])

if __name__ == '__main__':
    lines = generateLines()
    lineOverflows, pageOverflows = findOverflows(lines = lines, warn = True)
    doc = createDocx(lines)
    path = docxToPdf(doc)
    openPdf(path)


